"""Document text extraction service for syllabus and other uploads."""

import io
import logging
from typing import Optional, Tuple

logger = logging.getLogger(__name__)

# Supported file extensions and their MIME types
SUPPORTED_EXTENSIONS = {
    '.pdf': 'application/pdf',
    '.doc': 'application/msword',
    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    '.txt': 'text/plain',
}

SUPPORTED_MIME_TYPES = set(SUPPORTED_EXTENSIONS.values())


def is_supported_document(filename: str, content_type: str) -> bool:
    """Check if a file is a supported document type."""
    # Check by extension
    lower_filename = filename.lower()
    for ext in SUPPORTED_EXTENSIONS:
        if lower_filename.endswith(ext):
            return True

    # Check by MIME type
    if content_type in SUPPORTED_MIME_TYPES:
        return True

    return False


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from a PDF file."""
    try:
        from PyPDF2 import PdfReader

        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)

        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"Failed to extract text from PDF: {e}")
        raise ValueError(f"Could not extract text from PDF: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from a Word document (.docx)."""
    try:
        from docx import Document

        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)

        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {e}")
        raise ValueError(f"Could not extract text from Word document: {str(e)}")


def extract_text_from_doc(file_content: bytes) -> str:
    """
    Extract text from an old Word document (.doc).
    Note: python-docx doesn't support .doc files, so we provide a fallback message.
    """
    # python-docx doesn't support the old .doc format
    # For full .doc support, we'd need antiword or similar
    raise ValueError(
        "Old Word format (.doc) is not fully supported. "
        "Please save your document as .docx format and try again."
    )


def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from a plain text file."""
    try:
        # Try UTF-8 first, then fall back to other encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']

        for encoding in encodings:
            try:
                return file_content.decode(encoding)
            except UnicodeDecodeError:
                continue

        # Last resort: decode with errors ignored
        return file_content.decode('utf-8', errors='ignore')
    except Exception as e:
        logger.error(f"Failed to extract text from TXT: {e}")
        raise ValueError(f"Could not read text file: {str(e)}")


def extract_text(file_content: bytes, filename: str, content_type: str) -> Tuple[str, Optional[str]]:
    """
    Extract text from a document file.

    Args:
        file_content: The raw bytes of the file
        filename: Original filename (used to determine type)
        content_type: MIME type of the file

    Returns:
        Tuple of (extracted_text, error_message)
        If successful, error_message is None
        If failed, extracted_text is empty and error_message contains the error
    """
    lower_filename = filename.lower()

    try:
        # Determine file type and extract
        if lower_filename.endswith('.pdf') or content_type == 'application/pdf':
            text = extract_text_from_pdf(file_content)
        elif lower_filename.endswith('.docx') or content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            text = extract_text_from_docx(file_content)
        elif lower_filename.endswith('.doc') or content_type == 'application/msword':
            text = extract_text_from_doc(file_content)
        elif lower_filename.endswith('.txt') or content_type == 'text/plain':
            text = extract_text_from_txt(file_content)
        else:
            return "", f"Unsupported file type: {filename}"

        # Clean up the text
        text = text.strip()

        if not text:
            return "", "The document appears to be empty or contains no extractable text."

        logger.info(f"Successfully extracted {len(text)} characters from {filename}")
        return text, None

    except ValueError as e:
        return "", str(e)
    except Exception as e:
        logger.error(f"Unexpected error extracting text from {filename}: {e}")
        return "", f"Failed to extract text: {str(e)}"


def get_supported_extensions_display() -> str:
    """Get a human-readable list of supported extensions."""
    return ", ".join(sorted(SUPPORTED_EXTENSIONS.keys()))
