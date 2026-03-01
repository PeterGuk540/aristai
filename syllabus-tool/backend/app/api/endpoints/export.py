from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from app.schemas.syllabus import SyllabusData
from docx import Document
from io import BytesIO
import json
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

router = APIRouter()

from datetime import datetime, timedelta

def calculate_date(week_num_str: str, start_date_str: str) -> str:
    if not start_date_str:
        return ""
    try:
        # Try to parse week number
        week_num = int(week_num_str)
    except ValueError:
        return ""
    
    try:
        # Parse start date (YYYY-MM-DD)
        start_date = datetime.strptime(start_date_str, "%Y-%m-%d")
        # Calculate date
        target_date = start_date + timedelta(weeks=week_num - 1)
        # Format as "MMM D, YYYY" (e.g. Aug 25, 2024)
        return target_date.strftime("%b %d, %Y")
    except ValueError:
        return ""

@router.post("/export/docx")
async def export_docx(data: SyllabusData):
    try:
        doc = Document()
        
        # Title
        doc.add_heading(data.course_info.title, 0)
        
        # Course Info
        doc.add_heading('Course Information', level=1)
        doc.add_paragraph(f"Code: {data.course_info.code}")
        doc.add_paragraph(f"Instructor: {data.course_info.instructor}")
        doc.add_paragraph(f"Semester: {data.course_info.semester}")
        if data.course_info.email: doc.add_paragraph(f"Email: {data.course_info.email}")
        if data.course_info.office_hours: doc.add_paragraph(f"Office Hours: {data.course_info.office_hours}")
        if data.course_info.format: doc.add_paragraph(f"Format: {data.course_info.format}")
        
        if data.course_info.description:
            doc.add_heading('Course Description', level=2)
            doc.add_paragraph(data.course_info.description)
            
        if data.course_info.prerequisites:
            doc.add_heading('Prerequisites', level=2)
            doc.add_paragraph(data.course_info.prerequisites)
            
        if data.course_info.materials:
            doc.add_heading('Materials', level=2)
            doc.add_paragraph(data.course_info.materials)
        
        # Learning Goals
        doc.add_heading('Learning Goals', level=1)
        for goal in data.learning_goals:
            doc.add_paragraph(goal.text, style='List Bullet')
            
        # Schedule
        doc.add_heading('Schedule', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Week'
        hdr_cells[1].text = 'Date'
        hdr_cells[2].text = 'Topic'
        hdr_cells[3].text = 'Assignment'
        
        prev_week = None
        prev_week_cell = None
        prev_date = None
        prev_date_cell = None
        
        for item in data.schedule:
            row_cells = table.add_row().cells
            
            # Calculate date if missing
            display_date = item.date
            if not display_date and data.startDate:
                display_date = calculate_date(item.week, data.startDate)
            
            # Handle Week column merging
            if item.week == prev_week and prev_week_cell:
                row_cells[0].text = ""
                prev_week_cell.merge(row_cells[0])
            else:
                row_cells[0].text = str(item.week)
                prev_week = item.week
                prev_week_cell = row_cells[0]
            
            # Handle Date column merging
            if display_date == prev_date and prev_date_cell:
                row_cells[1].text = ""
                prev_date_cell.merge(row_cells[1])
            else:
                row_cells[1].text = display_date
                prev_date = display_date
                prev_date_cell = row_cells[1]
                
            row_cells[2].text = item.topic
            row_cells[3].text = item.assignment
            
        # Policies
        doc.add_heading('Policies', level=1)
        
        doc.add_heading('Academic Integrity', level=2)
        doc.add_paragraph(data.policies.academic_integrity)
        
        doc.add_heading('Accessibility', level=2)
        doc.add_paragraph(data.policies.accessibility)
        
        doc.add_heading('Attendance', level=2)
        doc.add_paragraph(data.policies.attendance)
        
        if data.policies.grading:
            doc.add_heading('Grading', level=2)
            doc.add_paragraph(data.policies.grading)
            
        if data.policies.late_work:
            doc.add_heading('Late Work', level=2)
            doc.add_paragraph(data.policies.late_work)
            
        if data.policies.communication:
            doc.add_heading('Communication', level=2)
            doc.add_paragraph(data.policies.communication)
            
        if data.policies.technology:
            doc.add_heading('Technology', level=2)
            doc.add_paragraph(data.policies.technology)
            
        if data.policies.learning_resources:
            doc.add_heading('Learning Resources', level=2)
            doc.add_paragraph(data.policies.learning_resources)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        headers = {
            'Content-Disposition': 'attachment; filename="syllabus.docx"'
        }
        
        return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/export/pdf")
async def export_pdf(data: SyllabusData):
    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Title
        story.append(Paragraph(data.course_info.title, styles['Title']))
        story.append(Spacer(1, 12))

        # Course Info
        story.append(Paragraph('Course Information', styles['Heading1']))
        story.append(Paragraph(f"Code: {data.course_info.code}", styles['Normal']))
        story.append(Paragraph(f"Instructor: {data.course_info.instructor}", styles['Normal']))
        story.append(Paragraph(f"Semester: {data.course_info.semester}", styles['Normal']))
        if data.course_info.email: story.append(Paragraph(f"Email: {data.course_info.email}", styles['Normal']))
        if data.course_info.office_hours: story.append(Paragraph(f"Office Hours: {data.course_info.office_hours}", styles['Normal']))
        if data.course_info.format: story.append(Paragraph(f"Format: {data.course_info.format}", styles['Normal']))
        story.append(Spacer(1, 12))
        
        if data.course_info.description:
            story.append(Paragraph('Course Description', styles['Heading2']))
            story.append(Paragraph(data.course_info.description, styles['Normal']))
            story.append(Spacer(1, 6))

        if data.course_info.prerequisites:
            story.append(Paragraph('Prerequisites', styles['Heading2']))
            story.append(Paragraph(data.course_info.prerequisites, styles['Normal']))
            story.append(Spacer(1, 6))

        if data.course_info.materials:
            story.append(Paragraph('Materials', styles['Heading2']))
            story.append(Paragraph(data.course_info.materials, styles['Normal']))
            story.append(Spacer(1, 6))

        # Learning Goals
        story.append(Paragraph('Learning Goals', styles['Heading1']))
        for goal in data.learning_goals:
            story.append(Paragraph(f"• {goal.text}", styles['Normal']))
        story.append(Spacer(1, 12))

        # Schedule
        story.append(Paragraph('Schedule', styles['Heading1']))
        table_data = [['Week', 'Date', 'Topic', 'Assignment']]
        
        week_spans = []
        date_spans = []
        
        current_week_start = 1 # 1-based index, 0 is header
        current_week_val = None
        
        current_date_start = 1
        current_date_val = None
        
        for i, item in enumerate(data.schedule):
            row_idx = i + 1
            week_val = str(item.week)
            
            # Calculate date if missing
            display_date = item.date
            if not display_date and data.startDate:
                display_date = calculate_date(item.week, data.startDate)
            
            row_data = [week_val, display_date, item.topic, item.assignment]
            
            # Week Span Logic
            if week_val == current_week_val:
                row_data[0] = '' # Clear text for merged cell
            else:
                # Close previous span
                if current_week_val is not None and row_idx - current_week_start > 1:
                    week_spans.append(('SPAN', (0, current_week_start), (0, row_idx - 1)))
                current_week_start = row_idx
                current_week_val = week_val
            
            # Date Span Logic
            if display_date == current_date_val:
                row_data[1] = '' # Clear text for merged cell
            else:
                # Close previous span
                if current_date_val is not None and row_idx - current_date_start > 1:
                    date_spans.append(('SPAN', (1, current_date_start), (1, row_idx - 1)))
                current_date_start = row_idx
                current_date_val = display_date
                
            table_data.append(row_data)
        
        # Close last spans
        if current_week_val is not None and len(data.schedule) + 1 - current_week_start > 1:
             week_spans.append(('SPAN', (0, current_week_start), (0, len(data.schedule))))
             
        if current_date_val is not None and len(data.schedule) + 1 - current_date_start > 1:
             date_spans.append(('SPAN', (1, current_date_start), (1, len(data.schedule))))

        t = Table(table_data)
        table_style = [
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'), # Center vertically for merged cells
        ]
        table_style.extend(week_spans)
        table_style.extend(date_spans)
        
        t.setStyle(TableStyle(table_style))
        story.append(t)
        story.append(Spacer(1, 12))

        # Policies
        story.append(Paragraph('Policies', styles['Heading1']))
        
        story.append(Paragraph('Academic Integrity', styles['Heading2']))
        story.append(Paragraph(data.policies.academic_integrity, styles['Normal']))
        story.append(Spacer(1, 6))
        
        story.append(Paragraph('Accessibility', styles['Heading2']))
        story.append(Paragraph(data.policies.accessibility, styles['Normal']))
        story.append(Spacer(1, 6))
        
        story.append(Paragraph('Attendance', styles['Heading2']))
        story.append(Paragraph(data.policies.attendance, styles['Normal']))
        story.append(Spacer(1, 6))

        if data.policies.grading:
            story.append(Paragraph('Grading', styles['Heading2']))
            story.append(Paragraph(data.policies.grading, styles['Normal']))
            story.append(Spacer(1, 6))

        if data.policies.late_work:
            story.append(Paragraph('Late Work', styles['Heading2']))
            story.append(Paragraph(data.policies.late_work, styles['Normal']))
            story.append(Spacer(1, 6))

        if data.policies.communication:
            story.append(Paragraph('Communication', styles['Heading2']))
            story.append(Paragraph(data.policies.communication, styles['Normal']))
            story.append(Spacer(1, 6))

        if data.policies.technology:
            story.append(Paragraph('Technology', styles['Heading2']))
            story.append(Paragraph(data.policies.technology, styles['Normal']))
            story.append(Spacer(1, 6))

        if data.policies.learning_resources:
            story.append(Paragraph('Learning Resources', styles['Heading2']))
            story.append(Paragraph(data.policies.learning_resources, styles['Normal']))
            story.append(Spacer(1, 6))

        doc.build(story)
        buffer.seek(0)
        
        headers = {
            'Content-Disposition': 'attachment; filename="syllabus.pdf"'
        }
        
        return StreamingResponse(buffer, media_type="application/pdf", headers=headers)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/export/json")
async def export_json(data: SyllabusData):
    try:
        json_str = data.model_dump_json(indent=2)
        buffer = BytesIO(json_str.encode('utf-8'))
        
        headers = {
            'Content-Disposition': 'attachment; filename="syllabus.json"'
        }
        
        return StreamingResponse(buffer, media_type="application/json", headers=headers)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/export/md")
async def export_md(data: SyllabusData):
    try:
        md_lines = []
        
        # Title
        md_lines.append(f"# {data.course_info.title}")
        md_lines.append("")
        
        # Course Info
        md_lines.append("## Course Information")
        md_lines.append(f"- **Code:** {data.course_info.code}")
        md_lines.append(f"- **Instructor:** {data.course_info.instructor}")
        md_lines.append(f"- **Semester:** {data.course_info.semester}")
        if data.course_info.email: md_lines.append(f"- **Email:** {data.course_info.email}")
        if data.course_info.office_hours: md_lines.append(f"- **Office Hours:** {data.course_info.office_hours}")
        if data.course_info.format: md_lines.append(f"- **Format:** {data.course_info.format}")
        md_lines.append("")
        
        if data.course_info.description:
            md_lines.append("### Course Description")
            md_lines.append(data.course_info.description)
            md_lines.append("")
            
        if data.course_info.prerequisites:
            md_lines.append("### Prerequisites")
            md_lines.append(data.course_info.prerequisites)
            md_lines.append("")
            
        if data.course_info.materials:
            md_lines.append("### Materials")
            md_lines.append(data.course_info.materials)
            md_lines.append("")
        
        # Learning Goals
        md_lines.append("## Learning Goals")
        for goal in data.learning_goals:
            md_lines.append(f"- {goal.text}")
        md_lines.append("")
        
        # Schedule
        md_lines.append("## Schedule")
        md_lines.append("| Week | Date | Topic | Assignment |")
        md_lines.append("|------|------|-------|------------|")
        
        prev_week = None
        prev_date = None
        
        for item in data.schedule:
            # Calculate date if missing
            display_date = item.date
            if not display_date and data.startDate:
                display_date = calculate_date(item.week, data.startDate)
                
            week_display = item.week if item.week != prev_week else ""
            date_display = display_date if display_date != prev_date else ""
            
            md_lines.append(f"| {week_display} | {date_display} | {item.topic} | {item.assignment} |")
            
            prev_week = item.week
            prev_date = display_date
            
        md_lines.append("")
        
        # Policies
        md_lines.append("## Policies")
        
        md_lines.append("### Academic Integrity")
        md_lines.append(data.policies.academic_integrity)
        md_lines.append("")
        
        md_lines.append("### Accessibility")
        md_lines.append(data.policies.accessibility)
        md_lines.append("")
        
        md_lines.append("### Attendance")
        md_lines.append(data.policies.attendance)
        md_lines.append("")

        if data.policies.grading:
            md_lines.append("### Grading")
            md_lines.append(data.policies.grading)
            md_lines.append("")

        if data.policies.late_work:
            md_lines.append("### Late Work")
            md_lines.append(data.policies.late_work)
            md_lines.append("")

        if data.policies.communication:
            md_lines.append("### Communication")
            md_lines.append(data.policies.communication)
            md_lines.append("")

        if data.policies.technology:
            md_lines.append("### Technology")
            md_lines.append(data.policies.technology)
            md_lines.append("")

        if data.policies.learning_resources:
            md_lines.append("### Learning Resources")
            md_lines.append(data.policies.learning_resources)
            md_lines.append("")
        
        content = "\n".join(md_lines)
        buffer = BytesIO(content.encode('utf-8'))
        
        headers = {
            'Content-Disposition': 'attachment; filename="syllabus.md"'
        }
        
        return StreamingResponse(buffer, media_type="text/markdown", headers=headers)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
