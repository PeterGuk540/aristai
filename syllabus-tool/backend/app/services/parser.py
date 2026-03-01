import fitz  # PyMuPDF
import docx
import io
import json
import concurrent.futures

from langchain_core.messages import HumanMessage, SystemMessage

from app.core.config import settings
from app.schemas.extraction import ExtractedSyllabusData, ExtractedSuggestions, ExtractedValidation, ValidationIssue
from app.services.json_extraction import JsonExtractionError, extract_first_json_object
from app.services.llm_factory import invoke_llm

def extract_text_from_pdf(file_content: bytes) -> str:
    text = ""
    with fitz.open(stream=file_content, filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_docx(file_content: bytes) -> str:
    doc = docx.Document(io.BytesIO(file_content))
    full_text = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)
        
    # Extract tables
    # Note: This appends tables at the end, which might lose context order,
    # but ensures the text is available for extraction.
    for table in doc.tables:
        for row in table.rows:
            # Join cells with spaces to mimic visual layout for regex
            row_text = [cell.text.strip() for cell in row.cells]
            # Filter out empty cells to avoid excessive spaces
            row_text = [t for t in row_text if t]
            if row_text:
                full_text.append("   ".join(row_text))
            
    return "\n".join(full_text)

def parse_file(filename: str, content: bytes) -> str:
    if filename.lower().endswith(".pdf"):
        return extract_text_from_pdf(content)
    elif filename.lower().endswith(".docx"):
        return extract_text_from_docx(content)
    else:
        # Try to read as text for any other extension (txt, md, py, etc.)
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("latin-1", errors="ignore")

def extract_metadata(filename: str, text: str) -> dict:
    metadata = {
        "school": "Unknown School",
        "department": "Unknown Department",
        "subject": "Unknown Subject"
    }
    
    # Simple heuristics based on common keywords
    text_lower = text.lower()
    filename_lower = filename.lower()
    
    # School detection
    if "bowling green" in text_lower or "bgsu" in text_lower:
        metadata["school"] = "Bowling Green State University"
    
    # Department detection
    if "department of" in text_lower:
        import re
        match = re.search(r"department of ([\w\s]+)", text_lower)
        if match:
            metadata["department"] = match.group(1).strip().title().split('\n')[0].strip()
    elif "accounting" in text_lower or "acct" in filename_lower:
        metadata["department"] = "Accounting"
    elif "computer science" in text_lower or "cs" in filename_lower:
        metadata["department"] = "Computer Science"

    # Subject detection (Course Code)
    import re
    # Look for patterns like "ACCT 101", "CS 3000"
    match = re.search(r"([A-Z]{2,4})\s?(\d{3,4})", text)
    if match:
        metadata["subject"] = f"{match.group(1)} {match.group(2)}"
    
    return metadata

def extract_syllabus_data_regex(text: str) -> dict:
    import re
    data = {
        "course_info": {
            "title": "",
            "code": "",
            "instructor": "",
            "semester": "",
            "description": "",
            "prerequisites": "",
            "office_hours": "",
            "email": "",
            "format": "",
            "materials": "",
        },
        "learning_goals": [],
        "schedule": [],
        "policies": {
            "academic_integrity": "",
            "accessibility": "",
            "attendance": "",
            "grading": "",
            "late_work": "",
            "communication": "",
            "technology": "",
            "learning_resources": "",
        },
    }
    
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    # 1. Course Info Extraction
    for line in lines[:50]: # Check first 50 lines for header info
        line_strip = line.strip()
        if not line_strip:
            continue

        # Instructor
        if not data["course_info"]["instructor"]:
            # Look for "Instructor: Name" or "Professor: Name"
            # Allow for leading chars like "* Instructor:" or "1. Instructor:"
            # Allow colon or hyphen as separator
            match = re.search(r"(?:Instructor|Professor|Teacher|Faculty)\s*[:\-]\s*(.*)", line_strip, re.IGNORECASE)
            if match:
                candidate = match.group(1).strip()
                # Clean up candidate (remove trailing chars, etc)
                candidate = re.sub(r"[_\*]+$", "", candidate).strip()
                
                if len(candidate) > 1:
                     # If it's too long, try to split it. 
                     # e.g. "Dr. Smith (Office Hours: Mon 1-2)" -> "Dr. Smith"
                     # Added semicolon and comma to splitters
                     if len(candidate) > 30:
                         # Try splitting by common separators
                         split_match = re.split(r"(\(|\||;|Office|Email|Ph|Phone)", candidate, flags=re.IGNORECASE)
                         if split_match:
                             candidate = split_match[0].strip()
                     
                     if len(candidate) < 100: # Increased limit
                         data["course_info"]["instructor"] = candidate
            else:
                # Also check for lines starting with "Instructor Name" without colon, but be strict
                match = re.match(r"^(?:Instructor|Professor)\s+(?![:\-])(.*)", line_strip, re.IGNORECASE)
                if match:
                     candidate = match.group(1).strip()
                     if len(candidate) < 60 and len(candidate) > 1 and not re.search(r"\b(is|are|will|should|must)\b", candidate):
                         data["course_info"]["instructor"] = candidate

        # Semester
        if not data["course_info"]["semester"]:
            # Look for "Semester: Fall 2024"
            # Allow colon or hyphen as separator
            match = re.search(r"(?:Semester|Term|Session)\s*[:\-]\s*(.*)", line_strip, re.IGNORECASE)
            if match:
                candidate = match.group(1).strip()
                candidate = re.sub(r"[_\*]+$", "", candidate).strip()
                
                if len(candidate) > 1:
                    if len(candidate) > 60:
                         split_match = re.split(r"(\(|\|)", candidate)
                         if split_match:
                             candidate = split_match[0].strip()

                    if len(candidate) < 100:
                        data["course_info"]["semester"] = candidate
            else:
                 match = re.match(r"^(?:Semester|Term|Session)\s+(?![:\-])(.*)", line_strip, re.IGNORECASE)
                 if match:
                     candidate = match.group(1).strip()
                     if len(candidate) < 60 and len(candidate) > 1 and not re.search(r"\b(is|are|will|should|must)\b", candidate):
                         data["course_info"]["semester"] = candidate
            
            # Fallback: Look for "Fall 2024", "Spring 2025" etc. directly
            if not data["course_info"]["semester"]:
                match = re.search(r"\b(Fall|Spring|Summer|Winter)\s+20\d{2}\b", line_strip, re.IGNORECASE)
                if match:
                    data["course_info"]["semester"] = match.group(0)

        # Course Title
        if not data["course_info"]["title"]:
             match = re.search(r"(?:Course Title|Title|Course Name)\s*[:]\s*(.*)", line_strip, re.IGNORECASE)
             if match:
                 candidate = match.group(1).strip()
                 if len(candidate) < 100 and len(candidate) > 1:
                     data["course_info"]["title"] = candidate
             # Heuristic: If line is all caps and long enough, and we are in the first 10 lines, it might be the title
             elif line_strip.isupper() and len(line_strip) > 10 and len(line_strip) < 100 and lines.index(line) < 10:
                 # Ignore common headers
                 if not any(x in line_strip for x in ["SYLLABUS", "COURSE", "UNIVERSITY", "DEPARTMENT", "SCHOOL"]):
                     data["course_info"]["title"] = line_strip

        # Course Code
        if not data["course_info"]["code"]:
             # Look for patterns like "ACCT 101", "CS 3000"
             # But ensure the line is not a long sentence
             match = re.search(r"\b([A-Z]{2,4})\s?(\d{3,4})\b", line_strip)
             if match and len(line_strip) < 100:
                 # Filter out common false positives
                 code_candidate = match.group(1)
                 if code_candidate not in ["TEXT", "FALL", "SPRI", "SUMM", "YEAR", "DATE", "TIME", "ROOM", "HALL", "BLDG"]:
                     data["course_info"]["code"] = f"{match.group(1)} {match.group(2)}"
                 
                 # If title is still missing, try to extract it from the same line if it looks like "Code: Title"
                 if not data["course_info"]["title"]:
                     # Check if there is text after the code
                     parts = line_strip.split(match.group(0))
                     if len(parts) > 1:
                         candidate = parts[1].strip(" :.-")
                         if len(candidate) > 3 and len(candidate) < 100:
                             data["course_info"]["title"] = candidate

        # Email
        if not data["course_info"]["email"]:
            match = re.search(r"\b([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})\b", line_strip)
            if match:
                data["course_info"]["email"] = match.group(1)
        
        # Office Hours
        if not data["course_info"]["office_hours"]:
            match = re.search(r"(?:Office\s*Hours|Hours)\s*[:\-]\s*(.*)", line_strip, re.IGNORECASE)
            if match:
                data["course_info"]["office_hours"] = match.group(1).strip()

        # Prerequisites
        if not data["course_info"]["prerequisites"]:
            match = re.search(r"(?:Prerequisites|Pre-requisites)\s*[:\-]\s*(.*)", line_strip, re.IGNORECASE)
            if match:
                data["course_info"]["prerequisites"] = match.group(1).strip()

        # Description (Simple one-line check, usually description is a block)
        if not data["course_info"]["description"]:
            match = re.search(r"(?:Course\s*Description|Description)\s*[:\-]\s*(.*)", line_strip, re.IGNORECASE)
            if match and len(match.group(1).strip()) > 20:
                data["course_info"]["description"] = match.group(1).strip()

        # Materials
        if not data["course_info"]["materials"]:
            match = re.search(r"(?:Materials|Textbook|Required\s*Text|Resources)\s*[:\-]\s*(.*)", line_strip, re.IGNORECASE)
            if match:
                data["course_info"]["materials"] = match.group(1).strip()

        # Format
        if not data["course_info"]["format"]:
            match = re.search(r"(?:Format|Modality|Class\s*Location)\s*[:\-]\s*(.*)", line_strip, re.IGNORECASE)
            if match:
                data["course_info"]["format"] = match.group(1).strip()

    # 2. Learning Goals
    # Look for a section header
    goals_started = False
    for line in lines:
        # Expanded header check
        # Matches: "Goals", "Course Goals", "Learning Goals", "Objectives", "Student Learning Outcomes", "Aims", "Competencies"
        # Allow leading numbering/bullets/whitespace
        if re.search(r"^\s*[\d\.\*\-]*\s*(?:student|course|class)?\s*(?:learning)?\s*(?:goals|objectives|outcomes|competencies|aims)(?:\s*[:])?\s*$", line, re.IGNORECASE):
            goals_started = True
            continue
            
        if goals_started:
            # Stop conditions:
            # 1. Hit another known section header
            if re.match(r"^\s*(?:grading|schedule|calendar|policies|academic integrity|required texts|prerequisites|evaluation|assessments)\s*[:]?\s*$", line, re.IGNORECASE):
                goals_started = False
                continue
            
            if line.strip():
                # If it's a list item, definitely add it
                if re.match(r"^[\d\.\-\•\*\+]\s+", line.strip()) or re.match(r"^\([a-z0-9]+\)", line.strip()):
                    data["learning_goals"].append({"id": len(data["learning_goals"]) + 1, "text": line.strip()})
                else:
                     # Not a list item.
                     # Check if it looks like a property/header (e.g. "Note:")
                     if re.match(r"^[A-Z][A-Za-z\s]+:$", line.strip()) and len(line.strip()) < 40:
                         # Likely a new section or property we don't want as a goal
                         pass
                     else:
                         # Treat as a goal line (even if unbulleted, like "test3")
                         data["learning_goals"].append({"id": len(data["learning_goals"]) + 1, "text": line.strip()}) 

    # 3. Schedule
    schedule_started = False
    week_counter = 1
    # specific flag for the "Week | Date | Topic | Assignment" table structure
    table_structure_detected = False
    
    for line in lines:
        # Check for section headers
        if re.search(r"^\s*(?:tentative)?\s*(?:course|class|weekly)?\s*(?:schedule|calendar|outline|plan|topics)\s*[:]?\s*$", line, re.IGNORECASE):
            schedule_started = True
            continue
        
        # Check for table headers (strong signal)
        # Enhanced to detect the specific user case: Week, Dates, Topics, Assignments
        if re.search(r"^\s*(?:week|session|module)\s+(?:date|time)?\s*(?:topic|subject|content)", line, re.IGNORECASE) or \
           (re.search(r"week", line, re.IGNORECASE) and re.search(r"date", line, re.IGNORECASE) and re.search(r"topic", line, re.IGNORECASE)):
            schedule_started = True
            table_structure_detected = True
            # Don't continue, this line might contain data or be the header itself
        
        if schedule_started:
             if re.match(r"^\s*(?:grading|policies|academic integrity|required texts|support services|resources|appendix)\s*[:]?\s*$", line, re.IGNORECASE):
                 schedule_started = False
        
        if schedule_started and line.strip():
            # Try to match the specific table row format: WeekNum DateRange Rest
            # e.g. "1 8/25 - 8/29 Introduction..."
            # Group 1: Week (digits)
            # Group 2: Date (digits/slashes/hyphens) - Optional but likely present if table_structure_detected
            # Group 3: Rest
            
            # Regex for "WeekNum DateRange Topic..."
            # We allow date to be optional in the regex but if we detected the table structure, we expect it.
            # \d{1,2}[/-]\d{1,2} matches 8/25
            date_pattern = r"(?:(?:\d{1,2}[/-]\d{1,2}(?:\s*-\s*\d{1,2}[/-]\d{1,2})?)|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2})"
            
            # Match: Start with digit (Week), then space, then Date (optional), then space, then Rest
            # We use a more specific regex first
            match_with_date = re.match(r"^(\d+)\s+(" + date_pattern + r")\s+(.+)", line.strip(), re.IGNORECASE)
            
            # Match: Start with digit (Week), then space, then Rest (no date or date not matched)
            match_simple = re.match(r"^(\d+)\s+(.+)", line.strip())
            
            # Match special rows like "Final Exam" or "Midterm"
            match_special = re.match(r"^(Final\s*Exam|Midterm|Exam\s*\d+)\s+(.+)", line.strip(), re.IGNORECASE)
            
            if match_with_date:
                week_num = match_with_date.group(1)
                date_str = match_with_date.group(2).strip()
                rest = match_with_date.group(3).strip()
                
                # Try to split rest into Topic and Assignment
                # If we have "Topics" and "Assignments" columns, they might be separated by multiple spaces
                parts = re.split(r"\s{2,}", rest)
                topic = parts[0]
                # If there are more parts, join them as assignment, or if just one more, it's assignment
                assignment = " ".join(parts[1:]) if len(parts) > 1 else ""
                
                data["schedule"].append({
                    "week": str(week_num),
                    "date": date_str,
                    "topic": topic,
                    "assignment": assignment
                })
                try:
                    week_counter = int(week_num) + 1
                except ValueError:
                    week_counter += 1
                
            elif match_simple:
                week_num = match_simple.group(1)
                rest = match_simple.group(2).strip()
                
                # Heuristic: If we are deep in the schedule (e.g. Week 10) and suddenly see Week 1,
                # it's likely a false positive from a numbered list in the footer/resources.
                try:
                    week_num_int = int(week_num)
                    if week_counter > 5 and week_num_int < week_counter - 4:
                        # Skip this likely false positive
                        continue
                except ValueError:
                    pass

                # Try to split rest into Topic and Assignment if there is a wide gap
                parts = re.split(r"\s{2,}", rest)
                topic = parts[0]
                assignment = parts[1] if len(parts) > 1 else ""
                
                data["schedule"].append({
                    "week": str(week_num),
                    "date": "",
                    "topic": topic,
                    "assignment": assignment
                })
                try:
                    week_counter = int(week_num) + 1
                except ValueError:
                    week_counter += 1
                
            elif match_special:
                # Handle special rows like "Final Exam"
                # They might not have a week number, so we use week_counter
                # They might have a date in the rest
                special_title = match_special.group(1).strip()
                rest = match_special.group(2).strip()
                
                # Try to find a date in the rest
                date_match = re.search(date_pattern, rest)
                date_str = ""
                if date_match:
                    date_str = date_match.group(0).strip()
                    # Remove date from rest to avoid duplication
                    rest = rest.replace(date_str, "").strip()
                
                # Try to split rest into Topic and Assignment
                parts = re.split(r"\s{2,}", rest)
                topic = parts[0]
                assignment = " ".join(parts[1:]) if len(parts) > 1 else ""
                
                # If topic is empty or just the special title repeated, use special title
                if not topic or topic.lower() == special_title.lower():
                    topic = special_title
                else:
                    topic = f"{special_title}: {topic}"
                
                # Use the special title as the week label if appropriate, or just the counter
                # User requested arbitrary text for week, so "Final Exam" is a valid week label.
                data["schedule"].append({
                    "week": special_title,
                    "date": date_str,
                    "topic": topic,
                    "assignment": assignment
                })
                week_counter += 1
                
            # Handle continuation lines (e.g. "Good" on next line)
            elif data["schedule"] and not re.match(r"^\s*(?:week|date|topic|assignment)", line, re.IGNORECASE):
                # Append to the last item's assignment if it exists, or topic
                last_item = data["schedule"][-1]
                # Heuristic: if the line is short or looks like a chapter ref, append to assignment
                # If it looks like text, append to topic?
                # In the user's table, "Topics" and "Assignments" are adjacent.
                # Let's append to assignment if we have one, otherwise topic.
                if last_item["assignment"]:
                    last_item["assignment"] += " " + line.strip()
                else:
                    # If assignment was empty, assume this line is the assignment
                    last_item["assignment"] = line.strip()
            
            # Simple heuristic: Look for "Week" or dates if we haven't matched a numbered line
            elif "week" in line.lower() or re.match(r"\d{1,2}/\d{1,2}", line):
                # Only add if we didn't just add a numbered week
                pass # We rely on the numbered week detection for this specific table format

    # 4. Policies
    # Extract blocks of text for policies
    current_policy = None
    for line in lines:
        lower_line = line.lower()
        if "academic integrity" in lower_line or "plagiarism" in lower_line:
            current_policy = "academic_integrity"
        elif "accessibility" in lower_line or "disability" in lower_line:
            current_policy = "accessibility"
        elif "attendance" in lower_line:
            current_policy = "attendance"
        elif "late work" in lower_line or "late policy" in lower_line or "late assignment" in lower_line:
            current_policy = "late_work"
        elif "communication" in lower_line or "email policy" in lower_line:
            current_policy = "communication"
        elif "technology" in lower_line or "technical requirements" in lower_line:
            current_policy = "technology"
        elif "learning resources" in lower_line or "student support" in lower_line or "tutoring" in lower_line:
            current_policy = "learning_resources"
        elif "grading" in lower_line:
            current_policy = "grading"
        elif "schedule" in lower_line:
            current_policy = None
        
        if current_policy and line.strip():
            data["policies"][current_policy] += line.strip() + " "

    return data

def extract_syllabus_data_llm(text: str, guidance_text: str = None, template_id: str = "BGSU_Standard") -> dict:
    def build_llm_context(full_text: str) -> str:
        # Keep both head and tail to reduce "missing policies/schedule" in long syllabi.
        max_chars = max(5000, int(getattr(settings, "LLM_MAX_CHARS", 30000)))
        head = full_text[: int(max_chars * 0.7)]
        tail = "" if len(full_text) <= len(head) else full_text[-int(max_chars * 0.3) :]

        # Add lightweight section hints (lines with common headers) to improve recall.
        import re

        lines = [ln.strip() for ln in full_text.splitlines() if ln.strip()]
        key_re = re.compile(
            r"(learning\s+goals|objectives|outcomes|competencies|schedule|calendar|outline|policies|attendance|academic\s+integrity|accessibility)",
            re.IGNORECASE,
        )
        hints = []
        for ln in lines:
            if key_re.search(ln):
                hints.append(ln)
            if len(hints) >= 40:
                break

        hint_block = "\n".join(hints)
        return (
            "[SECTION_HINTS]\n"
            + hint_block
            + "\n\n[TEXT_HEAD]\n"
            + head
            + ("\n\n[TEXT_TAIL]\n" + tail if tail else "")
        )

    def normalize_llm_output(data: ExtractedSyllabusData) -> dict:
        # Ensure ids/weeks are present and stable.
        goals = []
        for idx, g in enumerate(data.learning_goals, start=1):
            text_val = (g.text or "").strip()
            if not text_val:
                continue
            goals.append({"id": int(g.id) if g.id else idx, "text": text_val})

        schedule = []
        for idx, s in enumerate(data.schedule, start=1):
            topic_val = (s.topic or "").strip()
            if not topic_val:
                continue
            schedule.append(
                {
                    "week": str(s.week) if s.week else str(idx),
                    "date": (s.date or "").strip(),
                    "topic": topic_val,
                    "assignment": (s.assignment or "").strip(),
                }
            )

        result = {
            "course_info": {
                "title": (data.course_info.title or "").strip(),
                "code": (data.course_info.code or "").strip(),
                "instructor": (data.course_info.instructor or "").strip(),
                "semester": (data.course_info.semester or "").strip(),
                "description": (data.course_info.description or "").strip(),
                "prerequisites": (data.course_info.prerequisites or "").strip(),
                "office_hours": (data.course_info.office_hours or "").strip(),
                "email": (data.course_info.email or "").strip(),
                "format": (data.course_info.format or "").strip(),
                "materials": (data.course_info.materials or "").strip(),
            },
            "learning_goals": goals,
            "schedule": schedule,
            "policies": {
                "academic_integrity": (data.policies.academic_integrity or "").strip(),
                "accessibility": (data.policies.accessibility or "").strip(),
                "attendance": (data.policies.attendance or "").strip(),
                "grading": (data.policies.grading or "").strip(),
                "late_work": (data.policies.late_work or "").strip(),
                "communication": (data.policies.communication or "").strip(),
                "technology": (data.policies.technology or "").strip(),
                "learning_resources": (data.policies.learning_resources or "").strip(),
            },
        }
        
        if data.suggestions:
            result["suggestions"] = {
                "functional": data.suggestions.functional,
                "ui_ux": data.suggestions.ui_ux
            }
            
        if data.validation:
            result["validation"] = {
                "conforms_to_guidance": data.validation.conforms_to_guidance,
                "issues": [
                    i.model_dump() if hasattr(i, "model_dump") else i 
                    for i in data.validation.issues
                ]
            }
            
        return result

    context = build_llm_context(text)

    # Step 2: Generate Suggestions and Validation (If Guidance Exists)
    def run_validation(extracted_data: ExtractedSyllabusData, guidance: str) -> ExtractedSyllabusData:
        validation_system_prompt = (
            "You are a syllabus compliance auditor. You will be given the extracted syllabus data and a guidance document.\n"
            "Your task is to:\n"
            "1. Decompose the guidance document into a list of specific check items/requirements.\n"
            "2. For EACH check item, verify if the extracted content conforms to it.\n"
            "3. Return a list of validation results (both passed and failed checks). Do NOT omit passed checks.\n"
            "For each result, provide:\n"
            "   - status: 'passed' or 'failed'\n"
            "   - type: 'missing', 'incorrect', 'style' (if failed), or 'passed' (if passed)\n"
            "   - category: 'functional' or 'ui_ux'\n"
            "   - section: the section name (course_info, policies, schedule, learning_goals)\n"
            "   - field: the specific field name. For lists, include index/id (e.g. 'schedule.0.topic')\n"
            "   - issue: description of the requirement/check (e.g. 'Attendance policy must be included')\n"
            "   - current: a brief snippet of the current content (or 'Missing', or 'Present' if passed)\n"
            "   - suggestion: if failed, the specific content that should be there. If passed, use 'OK'.\n"
            "Output MUST be a single JSON object with 'validation' key matching the schema. Do NOT provide separate 'suggestions' list.\n"
            "Example Failed: {\"status\": \"failed\", \"type\": \"missing\", \"category\": \"functional\", \"section\": \"policies\", \"field\": \"attendance\", \"issue\": \"Attendance policy requirement\", \"current\": \"Missing\", \"suggestion\": \"Students are expected...\"}\n"
            "Example Passed: {\"status\": \"passed\", \"type\": \"passed\", \"category\": \"functional\", \"section\": \"policies\", \"field\": \"attendance\", \"issue\": \"Attendance policy requirement\", \"current\": \"Present\", \"suggestion\": \"OK\"}"
        )
        
        user_content = (
            f"EXTRACTED DATA:\n{extracted_data.model_dump_json()}\n\n"
            f"GUIDANCE DOCUMENT:\n{guidance}\n"
        )
        
        messages = [SystemMessage(content=validation_system_prompt), HumanMessage(content=user_content)]
        response = invoke_llm(messages)
        payload = extract_first_json_object(getattr(response, "content", ""))
        
        # Merge validation results back into extracted data
        if "suggestions" in payload:
            try:
                extracted_data.suggestions = ExtractedSuggestions(**payload["suggestions"])
            except Exception as e:
                print(f"Error parsing suggestions: {e}")
                
        if "validation" in payload:
            try:
                val_data = payload["validation"]
                if isinstance(val_data, list):
                    # Assume it's a list of issues
                    issues = []
                    for i in val_data:
                        if isinstance(i, dict):
                            issues.append(ValidationIssue(**i))
                        else:
                            issues.append(ValidationIssue(issue=str(i), type="unknown", section="unknown", field="unknown", current="unknown", suggestion="unknown"))
                    
                    extracted_data.validation = ExtractedValidation(
                        conforms_to_guidance=False,
                        issues=issues
                    )
                else:
                    extracted_data.validation = ExtractedValidation(**val_data)
            except Exception as e:
                print(f"Error parsing validation: {e}")
            
        return extracted_data

    # --- Execution ---
    # Combined System Prompt for Single-Pass Extraction
    if template_id == "Exponential":
        system_prompt = (
            "You are an expert syllabus parser specialized in 'Exponential Methodology' syllabi (Latin American context).\n"
            "Extract ALL details into a single JSON object matching the schema below.\n"
            "Schema:\n"
            "{\n"
            "  \"course_info\": {\"title\": \"\", \"code\": \"\", \"instructor\": \"\", \"semester\": \"\", \"description\": \"\", \"prerequisites\": \"\", \"office_hours\": \"\", \"email\": \"\", \"format\": \"\", \"materials\": \"\"},\n"
            "  \"learning_goals\": [{\"id\": 1, \"text\": \"...\"}],\n"
            "  \"schedule\": [{\"week\": \"1\", \"date\": \"...\", \"topic\": \"...\", \"assignment\": \"...\"}],\n"
            "  \"policies\": {\"academic_integrity\": \"\", \"accessibility\": \"\", \"attendance\": \"\", \"grading\": \"\", \"late_work\": \"\", \"communication\": \"\", \"technology\": \"\", \"learning_resources\": \"\"},\n"
            "  \"custom_sections\": {\n"
            "      \"sumilla\": \"\",\n"
            "      \"competencias\": \"\",\n"
            "      \"evaluacion_detallada\": \"\"\n"
            "  }\n"
            "}\n"
            "Rules:\n"
            "1. Course Info: Extract standard fields. Map 'Sumilla' to 'custom_sections.sumilla' AND 'course_info.description'.\n"
            "2. Learning Goals: Extract 'Resultados de Aprendizaje' into 'learning_goals'. Map 'Competencias' to 'custom_sections.competencias'.\n"
            "3. Schedule: Extract 'Programación de Contenidos'.\n"
            "4. Evaluation: Extract 'Evaluación del Aprendizaje' details into 'custom_sections.evaluacion_detallada'.\n"
            "Output ONLY the JSON object."
        )
    else:
        system_prompt = (
            "You are an expert syllabus parser. Extract ALL details into a single JSON object.\n"
            "Schema:\n"
            "{\n"
            "  \"course_info\": {\"title\": \"\", \"code\": \"\", \"instructor\": \"\", \"semester\": \"\", \"description\": \"\", \"prerequisites\": \"\", \"office_hours\": \"\", \"email\": \"\", \"format\": \"\", \"materials\": \"\"},\n"
            "  \"learning_goals\": [{\"id\": 1, \"text\": \"...\"}],\n"
            "  \"schedule\": [{\"week\": \"1\", \"date\": \"...\", \"topic\": \"...\", \"assignment\": \"...\"}],\n"
            "  \"policies\": {\"academic_integrity\": \"\", \"accessibility\": \"\", \"attendance\": \"\", \"grading\": \"\", \"late_work\": \"\", \"communication\": \"\", \"technology\": \"\", \"learning_resources\": \"\"}\n"
            "}\n"
            "Rules:\n"
            "1. Course Info: Extract all fields. Use empty string if unknown.\n"
            "2. Learning Goals: Extract all goals as a list.\n"
            "3. Schedule: Extract every week/session. Split distinct topics/assignments. Preserve dates.\n"
            "4. Policies: Extract full text for each policy category.\n"
            "Output ONLY the JSON object."
        )

    messages = [SystemMessage(content=system_prompt), HumanMessage(content=context)]
    
    try:
        response = invoke_llm(messages)
        data_dict = extract_first_json_object(getattr(response, "content", ""))
    except Exception as e:
        print(f"LLM Extraction failed: {e}")
        data_dict = {}

    # Construct full object from parts
    full_data = ExtractedSyllabusData(
        course_info=data_dict.get("course_info", {}),
        learning_goals=data_dict.get("learning_goals", []),
        schedule=data_dict.get("schedule", []),
        policies=data_dict.get("policies", {}),
        suggestions=None,
        validation=None
    )

    # Validation (if guidance exists)
    if guidance_text:
        try:
            full_data = run_validation(full_data, guidance_text)
        except Exception as e:
            print(f"LLM Validation failed: {e}")

    return normalize_llm_output(full_data)

def normalize_goal_text(text: str) -> str:
    import re
    # Remove leading numbers (1., 1), bullets (-, *, •), and parens ((1), (a))
    # Also lower case and strip
    clean = text.strip().lower()
    # Remove "1.", "1 ", "1-", etc.
    clean = re.sub(r"^[\d]+[\.\)\-]?\s*", "", clean)
    # Remove bullets
    clean = re.sub(r"^[\-\*\•\+]\s*", "", clean)
    # Remove (a), (1)
    clean = re.sub(r"^\([a-z0-9]+\)\s*", "", clean)
    return clean.strip()

def is_guidance_text(text: str) -> bool:
    """Check if the text looks like a guidance instruction rather than actual content."""
    if not text:
        return False
    text_lower = text.lower()
    indicators = [
        "must be", "should be", "full name of", "official title", 
        "term and year", "e.g.", "example:", "insert ", "provide "
    ]
    return any(indicator in text_lower for indicator in indicators)

def merge_syllabus_data(regex_data: dict, llm_data: dict) -> dict:
    if not llm_data:
        return regex_data
    
    # Deep copy to avoid modifying regex_data in place if we were to reuse it
    import copy
    merged = copy.deepcopy(regex_data)
    
    # Ensure basic structure exists in merged
    if "course_info" not in merged:
        merged["course_info"] = {}
    if "learning_goals" not in merged:
        merged["learning_goals"] = []
    if "schedule" not in merged:
        merged["schedule"] = []
    if "policies" not in merged:
        merged["policies"] = {}
    
    # 1. Course Info: Prefer LLM if available and looks valid, otherwise keep Regex
    for key in ["instructor", "semester", "title", "code", "description", "prerequisites", "office_hours", "email", "format", "materials"]:
        llm_val = llm_data.get("course_info", {}).get(key, "")
        regex_val = regex_data.get("course_info", {}).get(key, "")
        
        # Clean LLM value
        if isinstance(llm_val, str):
            llm_val = llm_val.strip()
            if llm_val.lower() in ["not found", "n/a", "unknown", "none", ""]:
                llm_val = None
            # Check if it looks like guidance text
            elif is_guidance_text(llm_val):
                llm_val = None
        
        if llm_val:
            # If LLM found something valid, use it.
            merged["course_info"][key] = llm_val
    
    # 2. Learning Goals: Union
    # Create a set of normalized texts to avoid duplicates
    existing_goals = {normalize_goal_text(g["text"]) for g in merged["learning_goals"]}
    next_id = len(merged["learning_goals"]) + 1
    
    for goal in llm_data.get("learning_goals", []):
        text = goal.get("text", "").strip()
        if text:
            normalized = normalize_goal_text(text)
            if normalized not in existing_goals:
                merged["learning_goals"].append({
                    "id": next_id,
                    "text": text
                })
                existing_goals.add(normalized)
                next_id += 1
            
    # 3. Schedule: Prefer LLM if available, UNLESS Regex found a high-quality schedule (many items) and LLM found few/none
    regex_schedule = regex_data.get("schedule", [])
    llm_schedule = llm_data.get("schedule", [])
    
    # If Regex found significantly more items than LLM (e.g. LLM missed the table), trust Regex
    if len(regex_schedule) > 5 and len(llm_schedule) < 3:
        merged["schedule"] = regex_schedule
    elif llm_schedule:
        merged["schedule"] = llm_schedule
    # Else keep regex_schedule (already in merged)
        
    # 4. Policies: Prefer LLM if available
    for key in ["academic_integrity", "accessibility", "attendance", "grading", "late_work", "communication", "technology", "learning_resources"]:
        llm_val = llm_data.get("policies", {}).get(key, "")
        if llm_val and str(llm_val).lower() not in ["not found", "n/a", "unknown", "none"]:
            merged["policies"][key] = llm_val
            
    # 5. Suggestions and Validation: Take from LLM
    if "suggestions" in llm_data:
        merged["suggestions"] = llm_data["suggestions"]
    if "validation" in llm_data:
        merged["validation"] = llm_data["validation"]
            
    return merged

def extract_syllabus_data(text: str, guidance_text: str = None, template_id: str = "BGSU_Standard") -> dict:
    import concurrent.futures
    
    # Run Regex and LLM extraction in parallel to save time
    with concurrent.futures.ThreadPoolExecutor() as executor:
        future_regex = executor.submit(extract_syllabus_data_regex, text)
        future_llm = executor.submit(extract_syllabus_data_llm, text, guidance_text, template_id)
        
        try:
            regex_data = future_regex.result()
        except Exception as e:
            print(f"Regex extraction failed: {e}")
            regex_data = {}
            
        try:
            llm_data = future_llm.result()
        except Exception as e:
            print(f"LLM extraction failed: {e}")
            llm_data = {}
    
    # 3. Merge Results
    final_data = merge_syllabus_data(regex_data, llm_data)
    
    return final_data

def merge_structured_data(data1: dict, data2: dict) -> dict:
    """
    Merges two structured syllabus data dictionaries.
    Prioritizes non-empty values from data2 over data1, but accumulates lists.
    """
    if not data1: return data2
    if not data2: return data1
    
    import copy
    merged = copy.deepcopy(data1)
    
    # 1. Course Info: Update if data2 has value
    for key in ["instructor", "semester", "title", "code", "description", "prerequisites", "office_hours", "email", "format", "materials"]:
        val2 = data2.get("course_info", {}).get(key, "")
        if val2:
            # Check if it looks like guidance text
            if not is_guidance_text(val2):
                merged["course_info"][key] = val2
            
    # 2. Learning Goals: Append unique
    existing_goals = {normalize_goal_text(g["text"]) for g in merged["learning_goals"]}
    next_id = len(merged["learning_goals"]) + 1
    
    learning_goals_2 = data2.get("learning_goals") or []
    for goal in learning_goals_2:
        text = goal.get("text", "").strip()
        if text:
            normalized = normalize_goal_text(text)
            if normalized not in existing_goals:
                merged["learning_goals"].append({
                    "id": next_id,
                    "text": text
                })
                existing_goals.add(normalized)
                next_id += 1
            
    # 3. Schedule: Append unique (by topic)
    # This is tricky. If data2 has a full schedule, maybe we should just take it?
    # Or append? Let's append for now, assuming different files cover different weeks.
    existing_topics = {s["topic"].strip().lower() for s in merged["schedule"]}
    next_week = len(merged["schedule"]) + 1
    
    schedule_2 = data2.get("schedule") or []
    for item in schedule_2:
        topic = item.get("topic", "").strip()
        if topic and topic.lower() not in existing_topics:
            # Adjust week number if needed, or keep original?
            # Let's just append.
            # If item has a week label, use it. If not, use next_week counter.
            week_label = item.get("week")
            if not week_label:
                week_label = str(next_week)
                next_week += 1
            else:
                week_label = str(week_label)
                # Try to increment next_week if the label is numeric, to keep counter in sync
                if week_label.isdigit():
                    next_week = int(week_label) + 1
                else:
                    next_week += 1

            merged["schedule"].append({
                "week": week_label,
                "date": item.get("date", ""),
                "topic": topic,
                "assignment": item.get("assignment", "")
            })
            existing_topics.add(topic.lower())
            
    # 4. Policies: Update if data2 has value (concatenate?)
    # Let's concatenate if both exist, to avoid losing info.
    for key in ["academic_integrity", "accessibility", "attendance", "grading", "late_work", "communication", "technology", "learning_resources"]:
        val1 = merged["policies"].get(key, "")
        val2 = data2.get("policies", {}).get(key, "")
        
        if val2:
            if val1 and val1.strip() != val2.strip():
                merged["policies"][key] = val1 + "\n\n" + val2
            else:
                merged["policies"][key] = val2
    
    # 5. Suggestions: Merge unique
    if "suggestions" in data2 and data2["suggestions"]:
        if "suggestions" not in merged or not merged["suggestions"]:
            merged["suggestions"] = {
                "functional": [],
                "ui_ux": []
            }
        
        # Merge functional
        existing_func = set(merged["suggestions"].get("functional", []) or [])
        for sugg in data2["suggestions"].get("functional", []) or []:
            if sugg not in existing_func:
                if "functional" not in merged["suggestions"] or merged["suggestions"]["functional"] is None:
                     merged["suggestions"]["functional"] = []
                merged["suggestions"]["functional"].append(sugg)
                existing_func.add(sugg)
                
        # Merge ui_ux
        existing_ui = set(merged["suggestions"].get("ui_ux", []) or [])
        for sugg in data2["suggestions"].get("ui_ux", []) or []:
            if sugg not in existing_ui:
                if "ui_ux" not in merged["suggestions"] or merged["suggestions"]["ui_ux"] is None:
                     merged["suggestions"]["ui_ux"] = []
                merged["suggestions"]["ui_ux"].append(sugg)
                existing_ui.add(sugg)

    # 6. Validation: Take the most recent non-empty validation
    if "validation" in data2 and data2["validation"]:
        if "validation" not in merged or not merged["validation"]:
            merged["validation"] = data2["validation"]
        else:
            # Merge issues
            # Issues are now dicts, so we deduplicate by 'issue' description
            existing_issues = (merged["validation"].get("issues", []) or [])
            existing_descs = set()
            
            # Handle both legacy strings and new dicts
            for i in existing_issues:
                if isinstance(i, dict):
                    existing_descs.add(i.get("issue", ""))
                elif isinstance(i, str):
                    existing_descs.add(i)

            new_issues = data2["validation"].get("issues", []) or []
            
            for issue in new_issues:
                desc = ""
                if isinstance(issue, dict):
                    desc = issue.get("issue", "")
                elif isinstance(issue, str):
                    desc = issue
                
                if desc and desc not in existing_descs:
                    if "issues" not in merged["validation"] or merged["validation"]["issues"] is None:
                        merged["validation"]["issues"] = []
                    merged["validation"]["issues"].append(issue)
                    existing_descs.add(desc)
            
            # Update conformance status
            if data2["validation"].get("conforms_to_guidance") is False:
                merged["validation"]["conforms_to_guidance"] = False

    return merged
