import re
import io
import docx


def extract_numbered_paragraphs(file_content: bytes) -> list[dict]:
    """Walk the DOCX in document order (body paragraphs + table cells) and produce a flat numbered list.

    Each dict includes table metadata (table_index, row_index, col_index) for table cells,
    with merged-cell deduplication via id(cell._tc).
    """
    doc = docx.Document(io.BytesIO(file_content))
    paragraphs = []
    index = 0

    # Body paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        para_type = "heading" if style_name.lower().startswith("heading") else "paragraph"
        paragraphs.append({
            "index": index,
            "text": para.text,
            "type": para_type,
            "location": "body",
            "table_index": None,
            "row_index": None,
            "col_index": None,
        })
        index += 1

    # Table cell paragraphs — deduplicate merged cells via id(cell._tc)
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            seen_tcs = set()
            for col_idx, cell in enumerate(row.cells):
                tc_id = id(cell._tc)
                if tc_id in seen_tcs:
                    continue
                seen_tcs.add(tc_id)
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    style_name = para.style.name if para.style else ""
                    para_type = "heading" if style_name.lower().startswith("heading") else "paragraph"
                    paragraphs.append({
                        "index": index,
                        "text": para.text,
                        "type": para_type,
                        "location": "table",
                        "table_index": table_idx,
                        "row_index": row_idx,
                        "col_index": col_idx,
                    })
                    index += 1

    return paragraphs


def group_paragraphs_by_section(paragraphs: list[dict]) -> dict:
    """Split the flat paragraph list into body + per-table groups.

    Returns:
        {
            "body": [paragraphs with location=="body"],
            "tables": [
                {
                    "table_index": 0,
                    "headers": ["Col1", "Col2", ...],
                    "paragraphs": [...],
                },
                ...
            ]
        }
    """
    body = [p for p in paragraphs if p["location"] == "body"]

    # Group table paragraphs by table_index
    table_map: dict[int, list[dict]] = {}
    for p in paragraphs:
        if p["location"] == "table" and p["table_index"] is not None:
            table_map.setdefault(p["table_index"], []).append(p)

    tables = []
    for t_idx in sorted(table_map.keys()):
        t_paras = table_map[t_idx]
        # Extract headers from row 0
        headers = []
        for p in t_paras:
            if p["row_index"] == 0:
                headers.append(p["text"].strip())

        tables.append({
            "table_index": t_idx,
            "headers": headers,
            "paragraphs": t_paras,
        })

    return {"body": body, "tables": tables}


def build_table_llm_prompt(table_group: dict, course_info: dict) -> str:
    """Build a grid-structured prompt for a single table.

    Uses [P#] markers so parse_llm_response works unchanged.
    Shows column headers and row structure for better LLM understanding.
    """
    headers = table_group["headers"]
    paras = table_group["paragraphs"]
    table_idx = table_group["table_index"]

    # Determine grid dimensions
    max_row = max((p["row_index"] for p in paras), default=0)
    max_col = max((p["col_index"] for p in paras), default=0)

    header_str = " | ".join(headers) if headers else "Unknown columns"

    lines = []
    lines.append(f"TABLE {table_idx} ({max_row + 1} rows x {max_col + 1} cols)")
    lines.append(f"Columns: {header_str}")
    lines.append("")

    # Group paragraphs by row
    row_map: dict[int, list[dict]] = {}
    for p in paras:
        row_map.setdefault(p["row_index"], []).append(p)

    for row_idx in sorted(row_map.keys()):
        row_paras = sorted(row_map[row_idx], key=lambda p: p["col_index"])
        row_label = "header" if row_idx == 0 else str(row_idx)
        for p in row_paras:
            col_name = headers[p["col_index"]] if p["col_index"] < len(headers) else f"Col {p['col_index']}"
            lines.append(f"[P{p['index']}] {p['text']}    (Row {row_label}, Col: {col_name})")

    return "\n".join(lines)


def build_llm_template_text(paragraphs: list[dict]) -> str:
    """Format the numbered paragraphs into a text block for the LLM."""
    lines = []
    for p in paragraphs:
        prefix = f"[P{p['index']}]"
        if p["type"] == "heading":
            prefix += " (heading)"
        if p["location"] == "table":
            prefix += " (table)"
        lines.append(f"{prefix} {p['text']}")
    return "\n".join(lines)


def parse_llm_response(response: str, total: int) -> dict[int, str]:
    """Parse the LLM's numbered output back into {paragraph_index: new_text}.

    Expected format: [P0] text here\n[P1] text here\n...
    If the LLM omits a paragraph number, that paragraph stays unchanged.
    """
    result = {}
    # Match [P<number>] followed by the rest of the line (greedy until next [P or end)
    # Use a pattern that captures multi-line content between markers
    pattern = re.compile(r'\[P(\d+)\]\s*(.*?)(?=\[P\d+\]|\Z)', re.DOTALL)
    for match in pattern.finditer(response):
        idx = int(match.group(1))
        text = match.group(2).strip().rstrip('|').strip()
        if idx < total:
            result[idx] = text
    return result


def fill_docx_by_index(file_content: bytes, paragraph_map: dict[int, str]) -> bytes:
    """Walk the original DOCX in the same order as extract_numbered_paragraphs.

    For each paragraph at index i, if i is in paragraph_map, replace the paragraph's
    text while preserving the first run's formatting.

    Uses merged-cell deduplication consistent with extract_numbered_paragraphs.
    """
    doc = docx.Document(io.BytesIO(file_content))
    index = 0

    # Body paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if index in paragraph_map:
            _set_paragraph_text(para, paragraph_map[index])
        index += 1

    # Table cell paragraphs — same dedup as extract_numbered_paragraphs
    for table in doc.tables:
        for row in table.rows:
            seen_tcs = set()
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id in seen_tcs:
                    continue
                seen_tcs.add(tc_id)
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    if index in paragraph_map:
                        _set_paragraph_text(para, paragraph_map[index])
                    index += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _set_paragraph_text(paragraph, new_text: str):
    """Replace a paragraph's text, preserving the first run's formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = new_text
