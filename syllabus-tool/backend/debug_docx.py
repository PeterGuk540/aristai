import docx
import io
import os

def extract_text_from_docx_current(file_path):
    with open(file_path, "rb") as f:
        file_content = f.read()
    doc = docx.Document(io.BytesIO(file_content))
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def extract_text_from_docx_improved(file_path):
    with open(file_path, "rb") as f:
        file_content = f.read()
    doc = docx.Document(io.BytesIO(file_content))
    
    full_text = []
    
    # Iterate over all elements in document body order would be ideal, but python-docx separates them.
    # A simple approach is to append tables at the end, or try to interleave.
    # But often tables are the main content for schedules.
    
    # Let's try to just get all text including tables.
    # Note: This doesn't preserve order perfectly if tables are interspersed with text, 
    # but it's better than missing them.
    # A better way is to iterate over `doc.element.body` but that's lower level.
    
    # For now, let's just append table text to see if it's there.
    
    for para in doc.paragraphs:
        full_text.append(para.text)
        
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append(" | ".join(row_text))
            
    return "\n".join(full_text)

file_path = "/data/syllabus_tool/backend/uploads/Yeh STAT 6010 5001 Fall 2025(final).docx_v1"
if os.path.exists(file_path):
    print("--- Current Extraction ---")
    print(extract_text_from_docx_current(file_path)[:500]) # Print first 500 chars
    print("...\n")
    
    print("--- Improved Extraction ---")
    improved = extract_text_from_docx_improved(file_path)
    print(improved[:500])
    print("...\n")
    
    if "Week" in improved and "Topic" in improved:
        print("FOUND 'Week' and 'Topic' in improved extraction!")
    else:
        print("Still didn't find keywords.")
else:
    print("File not found.")
